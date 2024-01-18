"""Automatic generation of reports for automatic analysis."""

from pathlib import Path
from itertools import combinations

import docx

from src.analysis import AutoAnalysis


class Reporter:
    """Reporter class generates docx reports for AutoAnalysis objects, allowing for easy sharing of results."""

    def __init__(
        self,
        analysis: AutoAnalysis,
        name: str = "Unnamed Report",
        output_path: Path = Path("out/report.docx"),
    ) -> None:
        """Initialise a Reporter object.

        Args:
            analysis (AutoAnalysis): Completed AutoAnalysis object.
            name (str, optional): Name of the report/experiment. Defaults to "Unnamed Report".
            output_path (Path, optional): Path to output report. Defaults to Path("out/report.docx").
        """
        self.report = docx.Document()
        self.name = name
        self.output_path = output_path
        self.analysis = analysis

    def generate_report(self) -> None:
        """Generate a full report for a given analysis."""
        self.report.add_heading(f"AutoAnalysis Report for {self.name}", 0)

        self._generate_premable()

        self._generate_data_description()

        if self.analysis.imputer:
            self._generate_imputation_description()

        if self.analysis.normaliser:
            self._generate_normalisation_description()

        if self.analysis.reducer:
            self._generate_reducer_description()

        if self.analysis.de_paths:
            self._generate_differential_expression_description()

        # Make sure the output directory exists
        self.output_path.parent.mkdir(parents=True, exist_ok=True)

        self.report.save(self.output_path)

    def _generate_premable(self) -> None:
        self.report.add_heading("About This Document", level=1)

        self.report.add_paragraph(
            "This document was automatically generated by AutoBioinformatics. All output files are available in the 'out' directory, and all plots are available in the 'img' directory."
        )

        self.report.add_paragraph(
            "For more information about AutoBioinformatics, please visit our website at https://github.com/stmball/auto_bioinformatics."
        )

    def _generate_data_description(self) -> None:
        self.report.add_heading("Data Description", level=1)

        num_groups = len(self.analysis.groups)
        self.report.add_paragraph(
            f"The data used in this analysis was split into {num_groups} groups: {', '.join(self.analysis.groups)}. "
        )

        self.report.add_paragraph(
            f"The data was made up of {self.analysis.data.shape[0]} genes and/or proteins. In total, {self.analysis.missing_data_percentage:.2f}% was missing."
        )

        self.report.add_paragraph(
            f"The data was scaled using the {self.analysis.scaler} scaler."
        )

        if self.analysis.scaler.explaination:
            self.report.add_paragraph(f"{self.analysis.scaler.explaination}")

    def _generate_imputation_description(self) -> None:
        self.report.add_heading("Imputation", level=1)

        self.report.add_paragraph(
            f"The data was imputed using {self.analysis.imputer}."
        )

        if self.analysis.imputer and self.analysis.imputer.explaination:
            self.report.add_paragraph(f"{self.analysis.imputer.explaination}")

        if self.analysis.imputation_plot_path:
            self.report.add_paragraph(
                "Below is a plot showing the distribution of the data before and after imputation."
            )

            self.report.add_picture(
                str(self.analysis.imputation_plot_path), width=docx.shared.Inches(6)
            )

    def _generate_normalisation_description(self) -> None:
        self.report.add_heading("Normalisation", level=1)

        self.report.add_paragraph(
            f"The data was normalised using {self.analysis.normaliser}."
        )

        if self.analysis.normaliser and self.analysis.normaliser.explaination:
            self.report.add_paragraph(f"{self.analysis.normaliser.explaination}")

    def _generate_reducer_description(self) -> None:
        self.report.add_heading("Dimensionality Reduction", level=1)

        self.report.add_paragraph(
            f"The data was reduced using {self.analysis.reducer}."
        )

        if self.analysis.reducer and self.analysis.reducer.explaination:
            self.report.add_paragraph(f"{self.analysis.reducer.explaination}")

        if self.analysis.dim_reducer_plot_path:
            self.report.add_paragraph(
                f"Below is a plot showing the {self.analysis.reducer} reduction of the data."
            )

            self.report.add_picture(
                str(self.analysis.dim_reducer_plot_path), width=docx.shared.Inches(6)
            )

    def _generate_differential_expression_description(self) -> None:
        self.report.add_heading("Differential Expression", level=1)

        self.report.add_paragraph(
            f"For each pair of groups, the fold change was calculated, and the p-value was calculated using a independent t-test."
        )

        self.report.add_paragraph(
            f"Samples were considered to be differentially expressed if the absolute fold change was greater than {self.analysis.log_fold_change_threshold} and the p-value was less than {self.analysis.p_value_threshold}."
        )

        self.report.add_paragraph(
            f"Below we will go through each pair of groups and discuss the differentially expressed genes."
        )

        for group_a, group_b in combinations(self.analysis.groups, 2):
            self.report.add_heading(
                f"Differential Expression Between {group_a} and {group_b}", level=2
            )

            self.report.add_paragraph(
                f"A volcano plot was generated for the comparison between {group_a} and {group_b}. Genes that were found to be differentially expressed are labelled."
            )

            if self.analysis.de_paths is not None:
                self.report.add_picture(
                    str(self.analysis.de_paths[f"{group_a}_{group_b}"]["volcano_fig"]),
                    width=docx.shared.Inches(6),
                )

            self.report.add_paragraph(
                f"Pathway analysis was performed on the differentially expressed genes, using the {self.analysis.gene_sets} database(s) for {self.analysis.organism}."
            )

            if self.analysis.de_paths is not None:
                self.report.add_picture(
                    str(self.analysis.de_paths[f"{group_a}_{group_b}"]["pathway_fig"]),
                    width=docx.shared.Inches(6),
                )
